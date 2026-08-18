import os
import time
import json
from pathlib import Path
from dotenv import load_dotenv
from groq import Groq
from pydantic import BaseModel, Field

load_dotenv()
my_api_key=os.getenv("GROQ_API_KEY")

if not my_api_key:
    raise ValueError("API key kaha hai bhai")



my_api_key = os.getenv("GROQ_API_KEY")
client = Groq(api_key=my_api_key)
model = "openai/gpt-oss-120b"

job_description="""
					Job Description 

Brief about the company:
InMobi:
Building a new company in the recession of 2007 was no ordinary task. Yet with passion and foresight, we charted our course, helping to transform the way consumers engage with their phones.
Over the last 17 years, InMobi has built a global Advertising Platform that powers our customers’ growth by helping them engage their audiences and drive real connections.
InMobi has also built a second unicorn, Glance, which is advancing digital consumption and creating a new wave of disruption. Present on 400M devices across India, SEA, Japan and the US – Glance is one of the largest content platforms globally with~200M daily active users.
Glance:
Founded in 2019, Glance is a consumer technology company that operates some of the most disruptive digital platforms including Glance, Roposo, and Glance TV. Glance has redefined the way the internet is consumed on the lock screen, removing the need to search for and download apps. Over 400 million smartphones now come enabled with Glance’s next-generation internet experience.     Roposo has revolutionized commerce by launching a destination for creator-led live entertainment commerce. Glance TV is changing the way consumers engage and interact with their televisions.     Headquartered in Singapore, Glance is an unconsolidated subsidiary of InMobi Group and is funded by Jio Platforms, Google, and Mithril Capital. For more information, visit glance.com, roposo.com, and inmobi.com.  


What will you be doing?
 
· Build the Future: Craft and launch cutting-edge features that bring our product roadmap to life.
· Code Like a Pro: Master new languages, write clean, high-quality code, and always push the boundaries of what’s possible.
· Full-Stack Wizardry: Be the go-to person for end-to-end support on our products and tools, keeping them running like a well-oiled machine.
· Design the Next Big Thing: Help shape and refine our tech processes, making everything we do smarter, faster, and better.

· Be the Voice: Share your knowledge through blogs and tech talks, and represent Glance as a thought leader in the tech world.
· Elevate the Team: Dive into code reviews, give valuable feedback, and help your peers level up their coding game.


What do we expect from you?
· Bachelor’s degree in computer science or equivalent.
· Good programming experience in Java/Python.
· Strong knowledge of Web Development Technologies and Cloud technologies.
· Should have decent knowledge on the cutting-edge technologies like AI/ML and prompt engineering.
· Well versed with the complete Software Engineering Lifecycle along with Agile practices.
· Strong Software development fundamentals including OOPs, MVC, general Algorithms, Design Patterns etc.
· Strong communication skills & writing skills
· Bonus: Significant programming experience for web applications, especially large-scale public-facing web sites
· Bonus: Significant programming experience in big data technologies.  


"""
class JobD(BaseModel):
    role: str
    required_skills: list[str]
    preferred_skills: list[str]
    minimum_experience: float | None
    education_requirements: list[str]
    responsibilities: list[str]

jobd_schema = JobD.model_json_schema()
def analyze_job_description(job_description):
    system_prompt = f"""
    You are an expert HR assistant.

    Your job is to analyze job descriptions and extract
    structured information from them.

    Return ONLY valid JSON matching this schema:

    {jobd_schema}

    IMPORTANT:
    Do NOT return the schema itself.
    Do NOT return fields like "properties", "title" or "type".
    Fill the schema with actual information extracted from the job description.

    If minimum experience is not mentioned, return null.
    If information for a list is missing, return an empty list.
    Do not invent information.
    """

    user_prompt = f"""
    Analyze the following job description:

    {job_description}
    """

    messages = [
        {
            "role": "system",
            "content": system_prompt
        },
        {
            "role": "user",
            "content": user_prompt
        }
    ]

    response = client.chat.completions.create(
        model=model,
        messages=messages,
        response_format={"type": "json_object"}
    )

    data = json.loads(response.choices[0].message.content)

    return JobD(**data)

system_prompt = f"""
You are an expert HR assistant.

Your job is to analyze job descriptions and extract
structured information from them.

Return ONLY valid JSON matching this schema:

{jobd_schema}
IMPORTANT:
Do NOT return the schema itself.
Do NOT return fields like "properties", "title" or "type".
Fill the schema with actual information extracted from the job description.

If minimum experience is not mentioned, return null.
If information for a list is missing, return an empty list.
Do not invent information.
"""

user_prompt = f"""
Analyze the following job description:

{job_description}
"""
message_system={
    "role" : "system",
    "content" : system_prompt
}
message_user={
    "role" : "user",
    "content" : user_prompt
}
response_format={
    "type" : "json_object"
}


messages=[message_system, message_user]

response=client.chat.completions.create(model=model, messages=messages, response_format=response_format)


answer=response.choices[0].message.content

raw_json=answer
# print(raw_json)




job_data=json.loads(raw_json)

job = JobD(**job_data)




#parse real
class MatchResult(BaseModel):
    score: float
    details: dict
class Experience(BaseModel):
    company: str | None = None
    role: str | None = None
    duration: str | None = None
    description: str | None = None
    skills_used: list[str] = []

class Education(BaseModel):
    degree: str | None = None
    institution: str | None = None
    grade: str | None = None
    year: int | str | None = None


class Project(BaseModel):
    name: str | None = None
    description: str | None = None


class Resume(BaseModel):
    name: str | None = None
    email: str | None = None
    phone: str | None = None

    total_experience_years: float | None = None

    skills: list[str] = []
    experiences: list[Experience] = []
    education: list[Education] = []
    projects: list[Project] = []
    certifications: list[str] = []

resume_schema = Resume.model_json_schema()
def final_score(job,resume):
    match_schema = MatchResult.model_json_schema()
    prompt = f"""
    You are an HR recruiter.

    Compare the candidate's resume with the job description.

    JOB DESCRIPTION:
    {job.model_dump_json(indent=2)}

    CANDIDATE RESUME:
    {resume.model_dump_json(indent=2)}
    Return JSON matching this schema:

    {match_schema}

    Give me:

    1. Candidate name
    2. Matching skills
    3. Missing important skills
    4. Whether experience requirement is met
    5. Overall match percentage from 0 to 100
    6. A short final verdict

    Keep the response concise and easy to read.
    """
    message={
        "role": "user",
        "content" : prompt
    }
    messages=[message]
    response_format={
        "type": "json_object"
    }
    response = client.chat.completions.create(model=model, messages=messages, response_format=response_format)
    data = json.loads(response.choices[0].message.content)
    return MatchResult(**data)
def parse_resume(resume_text):
    resume_text = resume_text[:25000]

    system_prompt = f"""
    You are an expert resume parser.

Extract information from the resume.

Return ONLY valid JSON with exactly these fields:

{{
    "name": null,
    "email": null,
    "phone": null,
    "total_experience_years": null,
    "skills": [],
    "experiences": [],
    "education": [],
    "projects": [],
    "certifications": []
}}

Rules:
- Do not invent information.
- Use null when a single value is unavailable.
- Use [] when a list has no information.
- Include internships inside experiences.
- Extract skills from the entire resume.
- Keep experience descriptions concise.
- Return ONLY JSON.
"""
    user_prompt = f"""
    Parse the following resume:

    {resume_text}
    """
    message_system={
        "role" : "system",
        "content" : system_prompt
    }
    message_user={
        "role" : "user",
        "content" : user_prompt
    }
    messages=[message_system, message_user]
    response_format={
        "type": "json_object"
    }
    response = client.chat.completions.create(
    model=model,
    messages=messages,
    response_format={"type": "json_object"},
    temperature=0
 )
    raw_output = response.choices[0].message.content
    data = json.loads(raw_output)
    resume = Resume(**data)
    return resume


from pypdf import PdfReader
from docx import Document
def read_pdf(file_path):
    reader = PdfReader(file_path)
    text = ""
    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            text += page_text + "\n"
    return text

def read_docx(file_path):
    document = Document(file_path)
    text = ""
    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            text += paragraph.text + "\n"
    
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text += cell.text + "\n"
    return text


def read_resume(file_path):
    if file_path.suffix.lower() == ".pdf":
        return read_pdf(file_path)
    elif file_path.suffix.lower() == ".docx":
        return read_docx(file_path)
    else:
        return None


def process_resume(file_path, job):
    resume_text = read_resume(Path(file_path))

    if not resume_text:
        return None

    parsed_resume = parse_resume(resume_text)
    result = final_score(job, parsed_resume)

    return {
        "name": parsed_resume.name or "Unknown Candidate",
        "score": result.score,
        "details": result.details
    }

